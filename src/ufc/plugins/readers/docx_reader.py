import re
from typing import Dict, Any, Iterable

from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl

from ufc.core.models import DocumentModel, ParagraphBlock, TableBlock
from ufc.plugins.registry import InputReader, PluginRegistry

_WS_RE = re.compile(r"[ \t]+")

def _clean_text(s: str) -> str:
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    lines = [_WS_RE.sub(" ", line).strip() for line in s.split("\n")]
    return "\n".join(lines).strip("\n")

def iter_block_items(doc: _Document) -> Iterable[object]:
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)

class DocxReader(InputReader):
    @classmethod
    def get_supported_extensions(cls) -> list[str]:
        return [".docx"]

    def read(self, file_path: str, options: Dict[str, Any]) -> DocumentModel:
        include_headers = options.get("include_headers", False)
        include_footers = options.get("include_footers", False)
        keep_empty = options.get("keep_empty_paragraphs", False)

        doc = Document(file_path)
        model = DocumentModel(metadata={"source": file_path, "type": "docx"})

        # Headers
        if include_headers:
            for i, sec in enumerate(doc.sections, start=1):
                model.blocks.append(ParagraphBlock(text=f"Section {i}", is_header=True))
                for p in sec.header.paragraphs:
                    t = _clean_text(p.text)
                    if t or keep_empty:
                        model.blocks.append(ParagraphBlock(text=t, is_header=True))
                
                for t in sec.header.tables:
                    model.blocks.append(self._parse_table(t, is_header=True))

        # Footers
        if include_footers:
            for i, sec in enumerate(doc.sections, start=1):
                model.blocks.append(ParagraphBlock(text=f"Section {i}", is_footer=True))
                for p in sec.footer.paragraphs:
                    t = _clean_text(p.text)
                    if t or keep_empty:
                        model.blocks.append(ParagraphBlock(text=t, is_footer=True))

                for t in sec.footer.tables:
                    model.blocks.append(self._parse_table(t, is_footer=True))

        # Body
        for block in iter_block_items(doc):
            if isinstance(block, Paragraph):
                t = _clean_text(block.text)
                if t or keep_empty:
                    model.blocks.append(ParagraphBlock(text=t))
            elif isinstance(block, Table):
                model.blocks.append(self._parse_table(block))

        return model

    def _parse_table(self, table: Table, is_header: bool = False, is_footer: bool = False) -> TableBlock:
        rows = []
        for row in table.rows:
            cells_text = [_clean_text(cell.text) for cell in row.cells]
            
            # Record effective cols if grid_cols_before/after are present
            before = int(getattr(row, "grid_cols_before", 0) or 0)
            after = int(getattr(row, "grid_cols_after", 0) or 0)
            
            padded = ([""] * before) + cells_text + ([""] * after)
            rows.append(padded)
            
        return TableBlock(rows=rows, is_header=is_header, is_footer=is_footer)

PluginRegistry.register_reader(DocxReader)
