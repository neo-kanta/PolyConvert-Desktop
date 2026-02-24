import os
import sys
import json
import re
import subprocess
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Tuple

import tkinter as tk
from tkinter import ttk, filedialog, messagebox

from docx import Document
from docx.document import Document as _Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl


# -----------------------------
# i18n (en-US, zh-TW)
# -----------------------------
I18N = {
    "en-US": {
        "title": "DOCX → TXT Converter (Batch + Chunk)",
        "language": "Language",
        "inputs": "Input .docx files",
        "add_files": "Add files…",
        "remove_selected": "Remove selected",
        "clear": "Clear",
        "output": "Output",
        "output_mode": "Output mode",
        "same_folder": "Same folder as each DOCX",
        "choose_folder": "Choose a folder",
        "browse_folder": "Browse folder…",
        "options": "Options",
        "include_tables": "Include tables",
        "table_format": "Table format",
        "tsv": "TSV (tab-separated)",
        "pipe": "Pipe table (| a | b |)",
        "normalize_tables": "Normalize irregular tables (best-effort)",
        "include_headers": "Include headers",
        "include_footers": "Include footers",
        "keep_empty": "Keep empty paragraphs",
        "utf8_bom": "Write UTF-8 with BOM (Windows-friendly)",
        "chunking": "Chunking (split output for upload)",
        "enable_chunk": "Enable chunk output",
        "chunk_size": "Chunk size (characters)",
        "overlap": "Overlap (characters)",
        "convert": "Convert",
        "open_output_folder": "Open output folder",
        "status_ready": "Ready.",
        "status_done": "Done ✅",
        "status_error": "Error ❌",
        "err_no_files": "Please add at least one .docx file.",
        "err_invalid_docx": "Some inputs are not .docx files.",
        "err_pick_outdir": "Please pick an output folder.",
        "ok_saved": "Saved:",
        "table_marker": "[TABLE]",
        "header_marker": "[HEADER]",
        "footer_marker": "[FOOTER]",
    },
    "zh-TW": {
        "title": "DOCX → TXT 轉換器（批次 + 分段）",
        "language": "語言",
        "inputs": "輸入 .docx 檔案",
        "add_files": "加入檔案…",
        "remove_selected": "移除選取",
        "clear": "清除",
        "output": "輸出",
        "output_mode": "輸出模式",
        "same_folder": "每個 DOCX 的同資料夾",
        "choose_folder": "指定輸出資料夾",
        "browse_folder": "瀏覽資料夾…",
        "options": "選項",
        "include_tables": "包含表格",
        "table_format": "表格格式",
        "tsv": "TSV（Tab 分隔）",
        "pipe": "直線表格（| a | b |）",
        "normalize_tables": "表格正規化（盡力處理合併/不規則）",
        "include_headers": "包含頁首",
        "include_footers": "包含頁尾",
        "keep_empty": "保留空白段落",
        "utf8_bom": "輸出 UTF-8 BOM（Windows 相容）",
        "chunking": "分段輸出（方便上傳）",
        "enable_chunk": "啟用分段輸出",
        "chunk_size": "每段大小（字元）",
        "overlap": "重疊（字元）",
        "convert": "開始轉換",
        "open_output_folder": "開啟輸出資料夾",
        "status_ready": "就緒。",
        "status_done": "完成 ✅",
        "status_error": "錯誤 ❌",
        "err_no_files": "請至少加入一個 .docx 檔案。",
        "err_invalid_docx": "有些輸入不是 .docx 檔。",
        "err_pick_outdir": "請選擇輸出資料夾。",
        "ok_saved": "已儲存：",
        "table_marker": "[表格]",
        "header_marker": "[頁首]",
        "footer_marker": "[頁尾]",
    },
}


CONFIG_PATH = Path.home() / ".docx_to_txt_gui_batch.json"


@dataclass
class AppConfig:
    lang: str = "en-US"
    last_dir: str = str(Path.home())
    out_dir: str = ""

    def save(self) -> None:
        try:
            CONFIG_PATH.write_text(json.dumps(self.__dict__, ensure_ascii=False, indent=2), encoding="utf-8")
        except Exception:
            pass

    @staticmethod
    def load() -> "AppConfig":
        try:
            data = json.loads(CONFIG_PATH.read_text(encoding="utf-8"))
            return AppConfig(
                lang=data.get("lang", "en-US"),
                last_dir=data.get("last_dir", str(Path.home())),
                out_dir=data.get("out_dir", ""),
            )
        except Exception:
            return AppConfig()


_WS_RE = re.compile(r"[ \t]+")


def _clean_text(s: str) -> str:
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    lines = [_WS_RE.sub(" ", line).strip() for line in s.split("\n")]
    return "\n".join(lines).strip("\n")


def iter_block_items(doc: _Document) -> Iterable[object]:
    body = doc.element.body
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            yield Table(child, doc)


def table_to_lines(table: Table, fmt: str, normalize: bool) -> List[str]:
    target_cols = 0
    if normalize:
        for row in table.rows:
            try:
                effective = int(getattr(row, "grid_cols_before", 0)) + len(row.cells) + int(getattr(row, "grid_cols_after", 0))
            except Exception:
                effective = len(row.cells)
            target_cols = max(target_cols, effective)

    lines: List[str] = []
    for row in table.rows:
        cells_text = [_clean_text(cell.text) for cell in row.cells]

        if normalize:
            before = int(getattr(row, "grid_cols_before", 0) or 0)
            after = int(getattr(row, "grid_cols_after", 0) or 0)
            padded = ([""] * before) + cells_text + ([""] * after)
            if target_cols > 0 and len(padded) < target_cols:
                padded += [""] * (target_cols - len(padded))
            cells_text = padded

        safe_cells = [c.replace("\n", "\\n") for c in cells_text]
        if fmt == "tsv":
            lines.append("\t".join(safe_cells).rstrip())
        else:
            lines.append("| " + " | ".join(safe_cells) + " |")

    return lines


def extract_docx(
    docx_path: str,
    include_tables: bool,
    table_format: str,
    normalize_tables: bool,
    include_headers: bool,
    include_footers: bool,
    keep_empty_paragraphs: bool,
    lang: str,
) -> str:
    doc = Document(docx_path)
    S = I18N[lang]
    out: List[str] = []

    if include_headers:
        for i, sec in enumerate(doc.sections, start=1):
            out.append(f"{S['header_marker']} Section {i}")
            for p in sec.header.paragraphs:
                t = _clean_text(p.text)
                if t or keep_empty_paragraphs:
                    out.append(t)
            if include_tables:
                for t in sec.header.tables:
                    out.append(S["table_marker"])
                    out.extend(table_to_lines(t, table_format, normalize_tables))
            out.append("")

    if include_footers:
        for i, sec in enumerate(doc.sections, start=1):
            out.append(f"{S['footer_marker']} Section {i}")
            for p in sec.footer.paragraphs:
                t = _clean_text(p.text)
                if t or keep_empty_paragraphs:
                    out.append(t)
            if include_tables:
                for t in sec.footer.tables:
                    out.append(S["table_marker"])
                    out.extend(table_to_lines(t, table_format, normalize_tables))
            out.append("")

    for block in iter_block_items(doc):
        if isinstance(block, Paragraph):
            t = _clean_text(block.text)
            if t or keep_empty_paragraphs:
                out.append(t)
        elif isinstance(block, Table):
            if include_tables:
                out.append(S["table_marker"])
                out.extend(table_to_lines(block, table_format, normalize_tables))
                out.append("")

    text = "\n".join(out).rstrip() + "\n"
    return text


def chunk_text(text: str, chunk_size: int, overlap: int) -> List[str]:
    if chunk_size <= 0:
        return [text]
    overlap = max(0, min(overlap, chunk_size - 1)) if chunk_size > 1 else 0

    chunks: List[str] = []
    n = len(text)
    start = 0
    while start < n:
        end = min(n, start + chunk_size)
        chunks.append(text[start:end])
        if end >= n:
            break
        start = end - overlap
    return chunks


def open_folder(folder: str) -> None:
    try:
        if sys.platform.startswith("win"):
            os.startfile(folder)  # type: ignore[attr-defined]
        elif sys.platform == "darwin":
            subprocess.run(["open", folder], check=False)
        else:
            subprocess.run(["xdg-open", folder], check=False)
    except Exception:
        pass


class DocxToTxtApp:
    def __init__(self, root: tk.Tk, cfg: AppConfig):
        self.root = root
        self.cfg = cfg

        self.lang = tk.StringVar(value=cfg.lang if cfg.lang in I18N else "en-US")
        self.files: List[str] = []

        self.output_mode = tk.StringVar(value="same")  # "same" or "folder"
        self.out_dir = tk.StringVar(value=cfg.out_dir)

        self.include_tables = tk.BooleanVar(value=True)
        self.table_format = tk.StringVar(value="tsv")
        self.normalize_tables = tk.BooleanVar(value=True)
        self.include_headers = tk.BooleanVar(value=False)
        self.include_footers = tk.BooleanVar(value=False)
        self.keep_empty = tk.BooleanVar(value=False)
        self.utf8_bom = tk.BooleanVar(value=False)

        self.enable_chunk = tk.BooleanVar(value=True)
        self.chunk_size = tk.IntVar(value=12000)  # chars
        self.overlap = tk.IntVar(value=300)       # chars

        self.status = tk.StringVar(value="")
        self.progress = tk.IntVar(value=0)

        self._build_ui()
        self._apply_i18n()

    def S(self, key: str) -> str:
        return I18N[self.lang.get()][key]

    def _build_ui(self) -> None:
        self.root.title("DOCX → TXT Converter")
        self.root.geometry("820x520")
        self.root.minsize(760, 480)

        main = ttk.Frame(self.root, padding=12)
        main.pack(fill="both", expand=True)

        # Language row
        top = ttk.Frame(main)
        top.pack(fill="x")
        self.lang_label = ttk.Label(top, text="Language")
        self.lang_label.pack(side="left")
        lang_box = ttk.Combobox(top, textvariable=self.lang, values=list(I18N.keys()), state="readonly", width=10)
        lang_box.pack(side="left", padx=8)
        lang_box.bind("<<ComboboxSelected>>", lambda e: self._apply_i18n())

        # Inputs
        inp = ttk.LabelFrame(main, text="Inputs")
        inp.pack(fill="both", expand=True, pady=10)

        left = ttk.Frame(inp)
        left.pack(side="left", fill="both", expand=True, padx=(8, 6), pady=8)

        self.files_list = tk.Listbox(left, height=10, selectmode=tk.EXTENDED)
        self.files_list.pack(fill="both", expand=True)

        right = ttk.Frame(inp)
        right.pack(side="right", fill="y", padx=(6, 8), pady=8)

        self.add_btn = ttk.Button(right, text="Add files…", command=self.add_files)
        self.add_btn.pack(fill="x", pady=(0, 6))

        self.remove_btn = ttk.Button(right, text="Remove selected", command=self.remove_selected)
        self.remove_btn.pack(fill="x", pady=6)

        self.clear_btn = ttk.Button(right, text="Clear", command=self.clear_files)
        self.clear_btn.pack(fill="x", pady=6)

        # Output
        out = ttk.LabelFrame(main, text="Output")
        out.pack(fill="x")

        mode_row = ttk.Frame(out)
        mode_row.pack(fill="x", padx=8, pady=(8, 4))

        self.same_radio = ttk.Radiobutton(mode_row, variable=self.output_mode, value="same", text="Same folder", command=self._toggle_outdir)
        self.same_radio.pack(side="left")

        self.folder_radio = ttk.Radiobutton(mode_row, variable=self.output_mode, value="folder", text="Choose a folder", command=self._toggle_outdir)
        self.folder_radio.pack(side="left", padx=12)

        dir_row = ttk.Frame(out)
        dir_row.pack(fill="x", padx=8, pady=(0, 8))

        self.dir_entry = ttk.Entry(dir_row, textvariable=self.out_dir)
        self.dir_entry.pack(side="left", fill="x", expand=True)

        self.dir_btn = ttk.Button(dir_row, text="Browse folder…", command=self.pick_outdir)
        self.dir_btn.pack(side="left", padx=8)

        # Options
        opt = ttk.LabelFrame(main, text="Options")
        opt.pack(fill="x", pady=10)

        self.tbl_chk = ttk.Checkbutton(opt, text="Include tables", variable=self.include_tables)
        self.tbl_chk.grid(row=0, column=0, sticky="w", padx=8, pady=6)

        self.fmt_label = ttk.Label(opt, text="Table format")
        self.fmt_label.grid(row=1, column=0, sticky="w", padx=8, pady=6)

        fmt = ttk.Frame(opt)
        fmt.grid(row=1, column=1, sticky="w", padx=8, pady=6)
        self.fmt_tsv = ttk.Radiobutton(fmt, text="TSV", value="tsv", variable=self.table_format)
        self.fmt_pipe = ttk.Radiobutton(fmt, text="Pipe", value="pipe", variable=self.table_format)
        self.fmt_tsv.pack(side="left")
        self.fmt_pipe.pack(side="left", padx=12)

        self.norm_chk = ttk.Checkbutton(opt, text="Normalize irregular tables", variable=self.normalize_tables)
        self.norm_chk.grid(row=2, column=0, sticky="w", padx=8, pady=6)

        self.hdr_chk = ttk.Checkbutton(opt, text="Include headers", variable=self.include_headers)
        self.hdr_chk.grid(row=3, column=0, sticky="w", padx=8, pady=6)

        self.ftr_chk = ttk.Checkbutton(opt, text="Include footers", variable=self.include_footers)
        self.ftr_chk.grid(row=3, column=1, sticky="w", padx=8, pady=6)

        self.empty_chk = ttk.Checkbutton(opt, text="Keep empty paragraphs", variable=self.keep_empty)
        self.empty_chk.grid(row=4, column=0, sticky="w", padx=8, pady=6)

        self.bom_chk = ttk.Checkbutton(opt, text="Write UTF-8 with BOM", variable=self.utf8_bom)
        self.bom_chk.grid(row=4, column=1, sticky="w", padx=8, pady=6)

        # Chunking
        chunk = ttk.LabelFrame(main, text="Chunking")
        chunk.pack(fill="x", pady=(0, 10))

        self.chunk_chk = ttk.Checkbutton(chunk, text="Enable chunk output", variable=self.enable_chunk, command=self._toggle_chunk_inputs)
        self.chunk_chk.grid(row=0, column=0, sticky="w", padx=8, pady=6)

        self.chunk_size_label = ttk.Label(chunk, text="Chunk size (characters)")
        self.chunk_size_label.grid(row=1, column=0, sticky="w", padx=8, pady=6)
        self.chunk_size_entry = ttk.Entry(chunk, textvariable=self.chunk_size, width=12)
        self.chunk_size_entry.grid(row=1, column=1, sticky="w", padx=8, pady=6)

        self.overlap_label = ttk.Label(chunk, text="Overlap (characters)")
        self.overlap_label.grid(row=2, column=0, sticky="w", padx=8, pady=6)
        self.overlap_entry = ttk.Entry(chunk, textvariable=self.overlap, width=12)
        self.overlap_entry.grid(row=2, column=1, sticky="w", padx=8, pady=6)

        # Actions / status
        actions = ttk.Frame(main)
        actions.pack(fill="x")

        self.convert_btn = ttk.Button(actions, text="Convert", command=self.convert_all)
        self.convert_btn.pack(side="left")

        self.open_btn = ttk.Button(actions, text="Open output folder", command=self.open_current_output)
        self.open_btn.pack(side="left", padx=10)

        self.pb = ttk.Progressbar(actions, variable=self.progress, maximum=100)
        self.pb.pack(side="right", fill="x", expand=True, padx=(10, 0))

        self.status_label = ttk.Label(main, textvariable=self.status)
        self.status_label.pack(anchor="e")

        self._toggle_outdir()
        self._toggle_chunk_inputs()

    def _apply_i18n(self) -> None:
        self.cfg.lang = self.lang.get()
        self.cfg.save()

        self.root.title(self.S("title"))

        self.lang_label.config(text=self.S("language"))
        self.add_btn.config(text=self.S("add_files"))
        self.remove_btn.config(text=self.S("remove_selected"))
        self.clear_btn.config(text=self.S("clear"))

        self.same_radio.config(text=self.S("same_folder"))
        self.folder_radio.config(text=self.S("choose_folder"))
        self.dir_btn.config(text=self.S("browse_folder"))

        self.tbl_chk.config(text=self.S("include_tables"))
        self.fmt_label.config(text=self.S("table_format"))
        self.fmt_tsv.config(text=self.S("tsv"))
        self.fmt_pipe.config(text=self.S("pipe"))
        self.norm_chk.config(text=self.S("normalize_tables"))
        self.hdr_chk.config(text=self.S("include_headers"))
        self.ftr_chk.config(text=self.S("include_footers"))
        self.empty_chk.config(text=self.S("keep_empty"))
        self.bom_chk.config(text=self.S("utf8_bom"))

        self.chunk_chk.config(text=self.S("enable_chunk"))
        self.chunk_size_label.config(text=self.S("chunk_size"))
        self.overlap_label.config(text=self.S("overlap"))

        self.convert_btn.config(text=self.S("convert"))
        self.open_btn.config(text=self.S("open_output_folder"))

        self.status.set(self.S("status_ready"))
        self.progress.set(0)

    def _toggle_outdir(self) -> None:
        is_folder = (self.output_mode.get() == "folder")
        state = "normal" if is_folder else "disabled"
        self.dir_entry.config(state=state)
        self.dir_btn.config(state=state)

    def _toggle_chunk_inputs(self) -> None:
        st = "normal" if self.enable_chunk.get() else "disabled"
        self.chunk_size_entry.config(state=st)
        self.overlap_entry.config(state=st)

    def add_files(self) -> None:
        paths = filedialog.askopenfilenames(
            title=self.S("inputs"),
            initialdir=self.cfg.last_dir,
            filetypes=[("Word Document", "*.docx"), ("All files", "*.*")],
        )
        if not paths:
            return
        self.cfg.last_dir = str(Path(paths[0]).parent)
        self.cfg.save()

        for p in paths:
            if p not in self.files:
                self.files.append(p)
                self.files_list.insert(tk.END, p)

    def remove_selected(self) -> None:
        sel = list(self.files_list.curselection())
        sel.reverse()
        for idx in sel:
            try:
                path = self.files_list.get(idx)
                if path in self.files:
                    self.files.remove(path)
                self.files_list.delete(idx)
            except Exception:
                pass

    def clear_files(self) -> None:
        self.files.clear()
        self.files_list.delete(0, tk.END)

    def pick_outdir(self) -> None:
        folder = filedialog.askdirectory(
            title=self.S("output"),
            initialdir=self.cfg.out_dir or self.cfg.last_dir,
        )
        if not folder:
            return
        self.out_dir.set(folder)
        self.cfg.out_dir = folder
        self.cfg.save()

    def _resolve_output_base(self, input_path: str) -> Path:
        if self.output_mode.get() == "same":
            return Path(input_path).parent
        outdir = self.out_dir.get().strip()
        if not outdir:
            raise ValueError(self.S("err_pick_outdir"))
        return Path(outdir)

    def convert_all(self) -> None:
        if not self.files:
            messagebox.showerror(self.S("status_error"), self.S("err_no_files"))
            return
        if any(Path(p).suffix.lower() != ".docx" for p in self.files):
            messagebox.showerror(self.S("status_error"), self.S("err_invalid_docx"))
            return

        total = len(self.files)
        self.progress.set(0)
        self.status.set("…")
        self.root.update_idletasks()

        encoding = "utf-8-sig" if self.utf8_bom.get() else "utf-8"
        saved_anywhere: Path | None = None

        chunk_root: Path | None = None
        if self.enable_chunk.get():
            if self.output_mode.get() == "folder":
                outdir = self.out_dir.get().strip()
                if not outdir:
                    messagebox.showerror(self.S("status_error"), self.S("err_pick_outdir"))
                    return
                base_for_all = Path(outdir)
            else:
                base_for_all = Path(self.files[0]).parent

            chunk_root = base_for_all / "ALL_CHUNKS"
            chunk_root.mkdir(parents=True, exist_ok=True)
            saved_anywhere = chunk_root

        try:
            for i, inp in enumerate(self.files, start=1):
                base = self._resolve_output_base(inp)
                stem = Path(inp).stem

                text = extract_docx(
                    docx_path=inp,
                    include_tables=self.include_tables.get(),
                    table_format=self.table_format.get(),
                    normalize_tables=self.normalize_tables.get(),
                    include_headers=self.include_headers.get(),
                    include_footers=self.include_footers.get(),
                    keep_empty_paragraphs=self.keep_empty.get(),
                    lang=self.lang.get(),
                )

                if self.enable_chunk.get():
                    size = int(self.chunk_size.get())
                    ov = int(self.overlap.get())
                    chunks = chunk_text(text, size, ov)

                    # ✅ write ALL chunks to ONE folder
                    assert chunk_root is not None
                    for idx, c in enumerate(chunks, start=1):
                        out_file = chunk_root / f"{stem}_part{idx:03d}.txt"
                        out_file.write_text(c, encoding=encoding, newline="\n")
                else:
                    out_file = base / f"{stem}.txt"
                    out_file.write_text(text, encoding=encoding, newline="\n")
                    saved_anywhere = base

                self.progress.set(int(i * 100 / total))
                self.status.set(f"{i}/{total}")
                self.root.update_idletasks()

            self.status.set(self.S("status_done"))
            if saved_anywhere:
                messagebox.showinfo(self.S("status_done"), f"{self.S('ok_saved')}\n{saved_anywhere}")

        except Exception as e:
            self.status.set(self.S("status_error"))
            messagebox.showerror(self.S("status_error"), f"{type(e).__name__}: {e}")

    def open_current_output(self) -> None:
        try:
            if self.output_mode.get() == "folder":
                outdir = self.out_dir.get().strip()
                if outdir:
                    open_folder(outdir)
            else:
                # same-folder mode: open folder of first file (best-effort)
                if self.files:
                    open_folder(str(Path(self.files[0]).parent))
        except Exception:
            pass


def main() -> None:
    cfg = AppConfig.load()
    root = tk.Tk()
    try:
        style = ttk.Style()
        if "clam" in style.theme_names():
            style.theme_use("clam")
    except Exception:
        pass
    DocxToTxtApp(root, cfg)
    root.mainloop()


if __name__ == "__main__":
    main()