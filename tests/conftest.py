import pytest
from docx import Document
from collections import namedtuple

@pytest.fixture
def sample_docx(tmp_path):
    doc = Document()
    doc.add_paragraph("Hello World")
    doc.add_paragraph("這是中文測試") # zh-TW
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "A1"
    table.cell(0, 1).text = "B1"
    table.cell(1, 0).text = "A2"
    table.cell(1, 1).text = "B2"
    
    # Adding a header
    sec = doc.sections[0]
    sec.header.paragraphs[0].text = "Header text"
    
    # Adding a footer
    sec.footer.paragraphs[0].text = "Footer text"

    path = tmp_path / "sample.docx"
    doc.save(str(path))
    return str(path)
